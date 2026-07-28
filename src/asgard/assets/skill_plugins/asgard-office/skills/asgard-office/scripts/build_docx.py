#!/usr/bin/env python3
"""Markdown document spec -> .docx.

    asgard skills run asgard-office -- build docx SPEC.md -o out.docx [--template NAME] [--values v.json]

The spec is Markdown with YAML front matter. Nothing in the pipeline needs Word,
LibreOffice, or pandoc — the same spec produces the same structure on every
machine, which is what makes a generated document reviewable instead of a
one-off artefact nobody can rebuild.
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx import Document  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK  # noqa: E402
from docx.shared import Mm, Pt, RGBColor  # noqa: E402

from officelib import compose, docxkit, mdblocks, specs  # noqa: E402

MAX_HEADING = 4  # Word's built-in heading styles thin out past this


def _rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


class DocxBuilder:
    def __init__(self, meta: specs.DocMeta, base: Path | None, assets: Path):
        self.meta = meta
        self.theme = meta.theme
        self.assets = assets
        self.document = Document(str(base)) if base else Document()
        self.counters = [0] * 6
        self.warnings: list[str] = []
        self._apply_page()
        self._apply_styles()

    # ------------------------------------------------------------- foundation

    def _apply_page(self) -> None:
        page = self.meta.page
        for section in self.document.sections:
            section.page_width = Mm(page.width_mm)
            section.page_height = Mm(page.height_mm)
            section.top_margin = Mm(page.top_mm)
            section.bottom_margin = Mm(page.bottom_mm)
            section.left_margin = Mm(page.left_mm)
            section.right_margin = Mm(page.right_mm)

    def _apply_styles(self) -> None:
        theme = self.theme
        cjk = theme.font_cjk or None
        normal = self.document.styles["Normal"]
        normal.font.size = Pt(theme.size_body)
        normal.paragraph_format.line_spacing = theme.line_spacing
        normal.paragraph_format.space_after = Pt(6)
        docxkit.set_style_fonts(normal, latin=theme.font_body, east_asian=cjk)
        sizes = {1: theme.size_h1, 2: theme.size_h2, 3: theme.size_h3, 4: theme.size_body + 0.5}
        for level, size in sizes.items():
            try:
                style = self.document.styles[f"Heading {level}"]
            except KeyError:  # a base template need not define every level
                continue
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = _rgb(theme.primary if level <= 2 else theme.secondary)
            style.paragraph_format.space_before = Pt(14 if level == 1 else 10)
            style.paragraph_format.space_after = Pt(4)
            docxkit.set_style_fonts(style, latin=theme.font_head, east_asian=cjk)

    @property
    def content_width_mm(self) -> float:
        page = self.meta.page
        return page.width_mm - page.left_mm - page.right_mm

    # ---------------------------------------------------------------- writing

    def _run(self, paragraph, item: mdblocks.Run, size: float | None, color: str | None):
        run = paragraph.add_run(item.text)
        run.bold = item.bold
        run.italic = item.italic
        if item.strike:
            run.font.strike = True
        face = self.theme.font_mono if item.code else self.theme.font_body
        docxkit.set_fonts(run, latin=face, east_asian=self.theme.font_cjk or None)
        run.font.size = Pt((size if size is not None else self.theme.size_body) - (0.5 if item.code else 0))
        if item.link:
            run.font.color.rgb = _rgb(self.theme.accent)
            run.font.underline = True
        elif color:
            run.font.color.rgb = _rgb(color)
        return run

    def add_runs(self, paragraph, runs, *, size: float | None = None, color: str | None = None):
        for item in runs:
            if not item.text:
                continue
            if item.link:
                # The run is built inside the paragraph, then relocated into the
                # hyperlink element — lxml moves rather than copies on append.
                docxkit.hyperlink(paragraph, item.link, lambda item=item: [self._run(paragraph, item, size, color)])
            else:
                self._run(paragraph, item, size, color)
        return paragraph

    def _line(self, text: str, *, size: float, color: str, bold: bool = False):
        paragraph = self.document.add_paragraph()
        self.add_runs(paragraph, [mdblocks.Run(text, bold=bold)], size=size, color=color)
        return paragraph

    # ----------------------------------------------------------------- fronts

    def cover(self) -> None:
        meta, theme = self.meta, self.theme
        self.document.add_paragraph().paragraph_format.space_after = Pt(150)
        if meta.title:
            self._line(meta.title, size=theme.size_title, color=theme.primary, bold=True)
        if meta.subtitle:
            self._line(meta.subtitle, size=theme.size_h2, color=theme.secondary)
        docxkit.paragraph_border(self.document.add_paragraph(), color=theme.accent, size=12)
        for value in (meta.company, meta.author, meta.date, f"Status: {meta.status}" if meta.status else ""):
            if value:
                self._line(value, size=theme.size_body, color=theme.muted)
        self.document.add_page_break()

    def title_block(self) -> None:
        meta, theme = self.meta, self.theme
        if meta.title:
            docxkit.keep_with_next(self._line(meta.title, size=theme.size_h1 + 6, color=theme.primary, bold=True))
        if meta.subtitle:
            self._line(meta.subtitle, size=theme.size_h2, color=theme.secondary)
        byline = " · ".join(part for part in (meta.author, meta.company, meta.date, meta.status) if part)
        if byline:
            self._line(byline, size=theme.size_caption, color=theme.muted)
        if meta.title or meta.subtitle or byline:
            docxkit.paragraph_border(self.document.add_paragraph(), color=theme.accent, size=8)

    def contents(self, label: str = "Contents") -> None:
        self._line(label, size=self.theme.size_h1, color=self.theme.primary, bold=True)
        docxkit.toc_field(self.document.add_paragraph())
        self.document.add_page_break()

    def running_marks(self) -> None:
        meta, theme = self.meta, self.theme
        for section in self.document.sections:
            if meta.header:
                paragraph = section.header.paragraphs[0]
                paragraph.text = ""
                self.add_runs(paragraph, [mdblocks.Run(meta.header)], size=theme.size_caption, color=theme.muted)
            footer = section.footer.paragraphs[0]
            footer.text = ""
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if meta.footer:
                self.add_runs(
                    footer, [mdblocks.Run(meta.footer + "   ")], size=theme.size_caption, color=theme.muted
                )
            docxkit.field_run(footer, "PAGE", placeholder="1")
            for run in footer.runs:
                run.font.size = Pt(theme.size_caption)
                run.font.color.rgb = _rgb(theme.muted)
                docxkit.set_fonts(run, latin=theme.font_body, east_asian=theme.font_cjk or None)

    # ------------------------------------------------------------------ blocks

    def heading(self, block: mdblocks.Block) -> None:
        level = min(max(block.level, 1), MAX_HEADING)
        try:
            paragraph = self.document.add_paragraph(style=f"Heading {level}")
        except KeyError:
            paragraph = self.document.add_paragraph()
        prefix = ""
        if self.meta.number_headings:
            self.counters[level - 1] += 1
            for deeper in range(level, len(self.counters)):
                self.counters[deeper] = 0
            prefix = ".".join(str(part) for part in self.counters[:level]) + ". "
        self.add_runs(paragraph, ([mdblocks.Run(prefix)] if prefix else []) + block.runs)
        docxkit.keep_with_next(paragraph)

    def bullets(self, block: mdblocks.Block) -> None:
        for item in mdblocks.normalise_levels(block.items):
            level = min(item.level, 2)
            family = "List Number" if item.ordered else "List Bullet"
            try:
                paragraph = self.document.add_paragraph(style=family if level == 0 else f"{family} {level + 1}")
            except KeyError:
                paragraph = self.document.add_paragraph()
                paragraph.paragraph_format.left_indent = Mm(6 + 6 * level)
            runs = list(item.runs)
            if item.checked is not None:
                runs.insert(0, mdblocks.Run("[x] " if item.checked else "[ ] "))
            self.add_runs(paragraph, runs)
            paragraph.paragraph_format.space_after = Pt(2)

    def table(self, block: mdblocks.Block) -> None:
        columns = max(len(block.header), max((len(row) for row in block.rows), default=0))
        if not columns:
            return
        table = self.document.add_table(rows=1, cols=columns)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        docxkit.table_borders(table, color=self.theme.muted)
        width = self.content_width_mm / columns
        for index, cell in enumerate(table.rows[0].cells):
            cell.width = Mm(width)
            runs = block.header[index] if index < len(block.header) else []
            self.add_runs(
                cell.paragraphs[0], [mdblocks.Run(mdblocks.plain(runs), bold=True)], color=self.theme.primary
            )
            docxkit.shade(cell._tc, self.theme.surface)
        docxkit.repeat_header(table.rows[0])
        for row in block.rows:
            cells = table.add_row().cells
            for index in range(columns):
                cells[index].width = Mm(width)
                self.add_runs(cells[index].paragraphs[0], row[index] if index < len(row) else [])
        self.document.add_paragraph().paragraph_format.space_after = Pt(2)

    def quote(self, block: mdblocks.Block) -> None:
        paragraph = self.document.add_paragraph()
        paragraph.paragraph_format.left_indent = Mm(8)
        docxkit.paragraph_border(paragraph, edge="left", color=self.theme.accent, size=18)
        self.add_runs(paragraph, block.runs, color=self.theme.secondary)

    def code(self, block: mdblocks.Block) -> None:
        for line in block.text.split("\n"):
            paragraph = self.document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.left_indent = Mm(4)
            docxkit.shade(paragraph._p, docxkit.MONO_SHADE)
            run = paragraph.add_run(line or " ")
            docxkit.set_fonts(run, latin=self.theme.font_mono)
            run.font.size = Pt(self.theme.size_body - 1)
        self.document.add_paragraph().paragraph_format.space_after = Pt(2)

    def image(self, block: mdblocks.Block) -> None:
        raw = Path(block.text)
        path = raw if raw.is_absolute() else (self.assets / raw)
        if not path.is_file():
            self.warnings.append(f"image not found, kept as a caption: {block.text}")
            self._line(f"[missing image: {block.text}]", size=self.theme.size_caption, color=self.theme.muted)
            return
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            paragraph.add_run().add_picture(str(path), width=Mm(min(self.content_width_mm, 160)))
        except Exception as exc:  # unreadable, or a format Word will not embed
            self.warnings.append(f"image could not be embedded ({exc.__class__.__name__}): {block.text}")
            return
        if block.alt:
            caption = self._line(block.alt, size=self.theme.size_caption, color=self.theme.muted)
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def emit(self, blocks: list[mdblocks.Block]) -> None:
        handlers = {
            "heading": self.heading,
            "para": lambda block: self.add_runs(self.document.add_paragraph(), block.runs),
            "list": self.bullets,
            "table": self.table,
            "quote": self.quote,
            "code": self.code,
            "image": self.image,
            "rule": lambda block: self.document.add_paragraph().add_run().add_break(WD_BREAK.PAGE),
        }
        for block in blocks:
            handler = handlers.get(block.kind)
            if handler:
                handler(block)

    def finish(self, out: Path) -> None:
        core = self.document.core_properties
        core.title = self.meta.title or core.title
        core.author = self.meta.author or core.author
        core.subject = self.meta.subtitle or core.subject
        out.parent.mkdir(parents=True, exist_ok=True)
        self.document.save(str(out))


# --------------------------------------------------------------------- driver


def build(
    spec_path: Path,
    out: Path,
    *,
    template_name: str = "",
    values_path: Path | None = None,
    root: Path | None = None,
) -> dict:
    front, body, template, missing = compose.resolve_spec(
        spec_path, template_name, values_path, root, ("docx", "md")
    )
    meta = specs.DocMeta.resolve(front)
    base = template.base if template and template.base and template.base.suffix in (".docx", ".dotx") else None
    blocks = mdblocks.parse(body)
    builder = DocxBuilder(meta, base, spec_path.resolve().parent)
    builder.cover() if meta.cover else builder.title_block()
    if meta.toc:
        builder.contents()
    builder.emit(blocks)
    builder.running_marks()
    builder.finish(out)
    return {
        "output": str(out),
        "template": template.name if template else None,
        "blocks": len(blocks),
        "headings": sum(1 for block in blocks if block.kind == "heading"),
        "unresolved_fields": missing,
        "warnings": builder.warnings,
    }


def main(argv: list[str]) -> int:
    parser = argparse.ArgumentParser(prog="build docx", description="Markdown document spec -> .docx")
    parser.add_argument("spec", type=Path)
    parser.add_argument("-o", "--output", type=Path, required=True)
    parser.add_argument("--template", default="", help="named template from the Office registry")
    parser.add_argument("--values", type=Path, default=None, help="JSON/YAML values for {{placeholders}}")
    parser.add_argument("--json", action="store_true", help="machine-readable build report")
    args = parser.parse_args(argv)
    try:
        report = build(args.spec, args.output, template_name=args.template, values_path=args.values)
    except (ValueError, OSError, KeyError) as exc:
        print(f"docx build failed: {exc}", file=sys.stderr)
        return 1
    if args.json:
        print(json.dumps(report, ensure_ascii=False, indent=2))
    else:
        print(f"wrote {report['output']}  ({report['blocks']} blocks, {report['headings']} headings)")
        for warning in report["warnings"]:
            print(f"  warning: {warning}")
        if report["unresolved_fields"]:
            print("  unresolved fields: " + ", ".join(report["unresolved_fields"]))
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
